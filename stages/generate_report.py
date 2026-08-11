#!/usr/bin/env python3
"""
OSINT Stage: Generate Report

Takes the company's llm_filtered.json (produced by llm_filter.py - already
fact-grounded, noise-filtered, and split into kept/excluded per category)
and renders it into a polished Word document for human review and
distribution.

DESIGN PRINCIPLE - the report body only ever shows "kept" data:
  - Every section in the report body is built exclusively from the "kept"
    side of llm_filtered.json. Nothing marked "excluded" by the LLM filter
    stage - and nothing that failed grounding - ever appears as a finding.
  - Everything that WAS excluded is not simply discarded: it's summarized
    in an "Appendix: Data Exclusions & Audit Log" section, with counts and
    per-item reasons, so a reviewer can audit what was left out and why
    without digging through raw JSON.
  - A section header is only printed if that section actually has kept
    content to show - no empty "Employees" heading followed by nothing.

Usage:
    python -m stages.generate_report --company "Resys Consultants"
"""

import argparse
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from lib.common import setup_logging, slugify_company
from lib.json_utils import load_json

logger = logging.getLogger(__name__)

# =====================================================================
# CONFIG / STYLE CONSTANTS
# =====================================================================

HEADER_FILL = "D9D9D9"  # light gray table header shading
NOTICE_FILL = "FCE8E8"  # pale red callout background
NOTICE_TEXT_COLOR = RGBColor(0x8A, 0x1F, 0x11)
ACCENT_COLOR = RGBColor(0x1F, 0x3B, 0x57)

TOOLS_USED = (
    "theHarvester, SpiderFoot, Amass, crt.sh/certificate-transparency, "
    "Metagoofil, Apify (search & breach lookup), Reacher (email validation), "
    "and a locally-hosted LLM (via Ollama) for noise filtering and "
    "summarization."
)

NOTICE_TEXT_BREACH = (
    "This section contains information about potential credential and "
    "personal data exposure sourced from third-party breach databases. It "
    "may include personally identifiable information (PII) about "
    "individuals associated with the target organization. Handle in "
    "accordance with your organization's data protection policy, restrict "
    "distribution to authorized personnel, and do not use this information "
    "for unauthorized access attempts of any kind."
)
NOTICE_TEXT_DARKWEB = (
    "This section contains information gathered from dark web / onion "
    "network sources. Findings may reference illicit marketplaces, leak "
    "sites, or forums. This information is provided for defensive security "
    "and awareness purposes only, may include PII, and should be handled "
    "under the same confidentiality restrictions as the breach exposure "
    "section above. Do not visit referenced dark web resources without "
    "appropriate authorization and safeguards."
)


# =====================================================================
# LOW-LEVEL DOCX HELPERS
# =====================================================================


def shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def shade_paragraph(paragraph, fill_hex: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_hyperlink(
    paragraph: Any,
    url: str,
    text: str,
    color: str = "0000FF",
    underline: bool = True,
) -> Any:
    """Adds a clickable hyperlink to a paragraph using underlying docx.oxml elements."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def add_table(doc: DocxDocument, headers: List[str], rows: List[List[Any]]) -> None:
    """Renders a standard docx table.

    Each cell value can be:
      - A string/number (rendered as text)
      - A (display_text, url) tuple or dict with 'url'/'text' (rendered as a clickable hyperlink)
      - None or empty string (rendered as '-')
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = ""
        run = header_cells[i].paragraphs[0].add_run(header)
        run.bold = True
        shade_cell(header_cells[i], HEADER_FILL)

    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = ""
            if isinstance(value, tuple) and len(value) == 2 and value[1]:
                text, url = value
                add_hyperlink(cells[i].paragraphs[0], str(url), str(text or url))
            elif isinstance(value, dict) and value.get("url"):
                url = str(value["url"])
                text = str(value.get("text") or url)
                add_hyperlink(cells[i].paragraphs[0], url, text)
            else:
                cells[i].paragraphs[0].text = "-" if value in (None, "") else str(value)

    doc.add_paragraph()  # spacing after table


def add_multiline_table(
    doc: DocxDocument, headers: List[str], rows: List[List[Any]]
) -> None:
    """Like add_table, but each cell value may be a list of strings, which
    are rendered as separate lines (paragraphs) within the cell instead of
    a single joined string - used for things like an employee's key facts."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        run = header_cells[i].paragraphs[0].add_run(header)
        run.bold = True
        shade_cell(header_cells[i], HEADER_FILL)

    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cell = cells[i]
            cell.paragraphs[0].text = ""
            if isinstance(value, list):
                lines = [str(v) for v in value if v]
                if not lines:
                    cell.paragraphs[0].text = "-"
                else:
                    cell.paragraphs[0].text = lines[0]
                    for line in lines[1:]:
                        cell.add_paragraph(line)
            else:
                cell.paragraphs[0].text = "-" if value in (None, "") else str(value)

    doc.add_paragraph()


def add_notice(doc: DocxDocument, text: str, label: str = "HANDLING NOTICE") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, NOTICE_FILL)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.color.rgb = NOTICE_TEXT_COLOR
    run_text = p.add_run(text)
    run_text.italic = True
    run_text.font.color.rgb = NOTICE_TEXT_COLOR


def add_section_heading(doc: DocxDocument, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ACCENT_COLOR


def format_timestamp(iso_str: Optional[str]) -> str:
    if not iso_str:
        return "unknown"
    try:
        dt = datetime.fromisoformat(iso_str)
        return dt.strftime("%B %d, %Y at %H:%M UTC")
    except ValueError:
        return iso_str


# =====================================================================
# TITLE PAGE & SUMMARY
# =====================================================================


def build_title_page(doc: DocxDocument, data: Dict[str, Any], report_date: str) -> None:
    title = doc.add_heading(data.get("company", "Unknown Company"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("OSINT Reconnaissance Report")
    run.font.size = Pt(16)
    run.italic = True
    run.font.color.rgb = ACCENT_COLOR

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Report generated: {report_date}").font.size = Pt(11)

    model_cfg = data.get("model", {}) or {}
    model_name = model_cfg.get("name", "unknown model")
    processed_at = format_timestamp(data.get("generated_at"))

    doc.add_paragraph()
    add_section_heading(doc, "Methodology", level=2)
    doc.add_paragraph(
        "This report was produced by an automated OSINT reconnaissance pipeline "
        "that discovers an organization's public-facing digital footprint - "
        "domains and infrastructure, employees, email addresses, publicly "
        "accessible documents, known credential/breach exposure, and dark web "
        "mentions - using a combination of open-source tools and third-party "
        "lookup services."
    )
    doc.add_paragraph(f"Tools used: {TOOLS_USED}")
    doc.add_paragraph(
        f"Raw findings were subsequently filtered and summarized by a locally-"
        f"hosted language model ({model_name}), processed on {processed_at}, "
        f"to remove duplicate and low-value noise (e.g. repeated automated "
        f"scan artifacts) and to produce plain-language summaries."
    )

    add_notice(
        doc,
        "This report's contents were partially generated and filtered by an "
        "AI language model. While the pipeline is designed to ground every "
        "AI-generated judgment in the underlying collected data and never "
        "invent new facts, AI-assisted output can still contain errors or "
        "misjudgments. This report must be reviewed by a qualified analyst "
        "before being relied upon, distributed, or acted upon.",
        label="HUMAN REVIEW REQUIRED",
    )

    doc.add_page_break()


def build_summary_table(doc: DocxDocument, data: Dict[str, Any]) -> None:
    stats = data.get("stats", {}) or {}
    if not stats:
        return

    add_section_heading(doc, "Summary of Findings", level=1)

    label_map = [
        ("domains_kept", "Domains identified"),
        ("emails_kept", "Email addresses identified"),
        ("employees_kept", "Employees / associated individuals identified"),
        ("breaches_kept", "Accounts with confirmed breach exposure"),
        ("darkweb_with_mentions", "Dark web targets with mentions found"),
        ("documents_summarized", "Public documents reviewed"),
    ]
    rows = [[label, stats[key]] for key, label in label_map if key in stats]
    if rows:
        add_table(doc, ["Category", "Count"], rows)


# =====================================================================
# INFRASTRUCTURE
# =====================================================================


def build_infrastructure_section(doc: DocxDocument, data: Dict[str, Any]) -> None:
    domains_kept = data.get("domains", {}).get("kept", [])
    dns_infra = data.get("dns_infra", {}) or {}
    dns_rows = [
        [
            host,
            ", ".join(str(r) for r in records)
            if isinstance(records, list)
            else str(records),
        ]
        for host, records in dns_infra.items()
        if records
    ]

    if not domains_kept and not dns_rows:
        return

    add_section_heading(doc, "Infrastructure", level=1)

    if domains_kept:
        add_section_heading(doc, "Domains", level=2)
        rows = []
        for d in domains_kept:
            note = (d.get("_llm_verdict") or {}).get("note", "")
            rows.append(
                [
                    d.get("domain", ""),
                    ", ".join(d.get("sources", [])),
                    "Yes" if d.get("dns_validated") else "No",
                    note,
                ]
            )
        add_table(doc, ["Domain", "Discovery Source(s)", "DNS Validated", "Note"], rows)

    if dns_rows:
        add_section_heading(doc, "DNS Infrastructure", level=2)
        add_table(doc, ["Host", "Records"], dns_rows)


# =====================================================================
# EMPLOYEES
# =====================================================================

EMPLOYEE_TIER_SORT_ORDER = {
    "leadership": 0,
    "current_employee": 1,
    "intern": 2,
    "former_employee": 3,
}
EMPLOYEE_TIER_LABELS = {
    "leadership": "Leadership",
    "current_employee": "Current Employee",
    "intern": "Intern",
    "former_employee": "Former Employee",
}


def _employee_tier(e: Dict[str, Any]) -> str:
    verdict = e.get("_llm_verdict") or {}
    return e.get("employee_tier") or verdict.get("tier") or ""


def _employee_tier_sort_key(e: Dict[str, Any]) -> int:
    return EMPLOYEE_TIER_SORT_ORDER.get(
        _employee_tier(e), len(EMPLOYEE_TIER_SORT_ORDER)
    )


def _employee_social_links(e: Dict[str, Any]) -> List[str]:
    lines: List[str] = []
    seen_urls = set()

    linkedin_url = e.get("linkedin_url")
    if linkedin_url and linkedin_url not in seen_urls:
        lines.append(f"LinkedIn: {linkedin_url}")
        seen_urls.add(linkedin_url)

    for entry in e.get("additional_profiles") or []:
        if not isinstance(entry, dict):
            continue
        for platform in entry.get("platforms") or []:
            if not isinstance(platform, dict):
                continue
            url = platform.get("url")
            if not url or url in seen_urls:
                continue
            site = platform.get("site") or "Profile"
            lines.append(f"{site}: {url}")
            seen_urls.add(url)

    return lines


def build_employees_section(doc: DocxDocument, data: Dict[str, Any]) -> None:
    employees_kept = data.get("employees", {}).get("kept", [])
    if not employees_kept:
        return

    add_section_heading(doc, "Employees", level=1)
    doc.add_paragraph(
        "Individuals below were identified as publicly associated with the "
        "target organization, primarily via professional networking profiles. "
        "Listed by confidence in their connection to the target company "
        "(leadership, then current employees, interns, and former employees)."
    )

    sorted_employees = sorted(employees_kept, key=_employee_tier_sort_key)

    rows = []
    for e in sorted_employees:
        verdict = e.get("_llm_verdict") or {}
        connection = verdict.get("connection_to_target") or "-"
        key_facts = verdict.get("key_facts") or []
        tier = _employee_tier(e)
        rows.append(
            [
                e.get("name", "-"),
                EMPLOYEE_TIER_LABELS.get(tier, tier or "-"),
                e.get("job_title") or "-",
                connection,
                key_facts,
                _employee_social_links(e) or ["-"],
            ]
        )
    add_multiline_table(
        doc,
        [
            "Name",
            "Tier",
            "Job Title",
            "Connection to Target",
            "Key Facts",
            "Social Links",
        ],
        rows,
    )


# =====================================================================
# EMAILS
# =====================================================================


def _email_domain(email: Optional[str]) -> str:
    if not email or "@" not in email:
        return ""
    return email.split("@", 1)[1].strip().lower()


def _pattern_preference_rank(email: str) -> int:
    local = email.split("@", 1)[0] if "@" in email else email
    if "." in local:
        return 0
    if "_" in local:
        return 1
    return 2


def build_emails_section(doc: DocxDocument, data: Dict[str, Any]) -> None:
    emails_kept = data.get("emails", {}).get("kept", [])
    if not emails_kept:
        return

    add_section_heading(doc, "Emails", level=1)

    email_domains = data.get("email_domains", []) or []
    catchall_domains = {
        d.get("domain") for d in email_domains if d.get("confirmed_catch_all")
    }

    verified_emails = [
        e for e in emails_kept if _email_domain(e.get("email")) not in catchall_domains
    ]
    catchall_emails = [
        e for e in emails_kept if _email_domain(e.get("email")) in catchall_domains
    ]

    if verified_emails:
        rows = []
        for e in verified_emails:
            verdict = e.get("_llm_verdict") or {}
            rows.append(
                [
                    e.get("email", ""),
                    e.get("employee") or "-",
                    e.get("validation_status", "-"),
                    verdict.get("tier", "-"),
                ]
            )
        add_table(
            doc,
            [
                "Email Address",
                "Associated Employee",
                "Validation Status",
                "Confidence Tier",
            ],
            rows,
        )

    domains_present = sorted({_email_domain(e.get("email")) for e in catchall_emails})
    for domain in domains_present:
        domain_emails = [
            e for e in catchall_emails if _email_domain(e.get("email")) == domain
        ]
        doc.add_paragraph(
            f"Email validation for {domain} was inconclusive - this mail server "
            f"accepts all addresses (catch-all configuration), so individual "
            f"mailbox existence cannot be confirmed via SMTP. "
            f"{len(domain_emails)} name-pattern-based candidates were generated "
            f"but are unverified."
        )

        best_per_employee: Dict[str, str] = {}
        for e in domain_emails:
            employee = e.get("employee") or "-"
            email = e.get("email") or ""
            current = best_per_employee.get(employee)
            if current is None or _pattern_preference_rank(
                email
            ) < _pattern_preference_rank(current):
                best_per_employee[employee] = email
        add_table(
            doc,
            ["Employee", "Most Likely Address"],
            [[name, addr] for name, addr in sorted(best_per_employee.items())],
        )


# =====================================================================
# DOCUMENTS
# =====================================================================


def build_documents_section(doc: DocxDocument, data: Dict[str, Any]) -> None:
    documents = data.get("documents", []) or []
    usable_docs = [d for d in documents if d.get("summary")]
    if not usable_docs:
        return

    add_section_heading(doc, "Documents", level=1)
    doc.add_paragraph(
        "Publicly accessible documents discovered on the target organization's "
        "domain(s), with a one-line usability summary for each."
    )

    rows = []
    for d in usable_docs:
        url = d.get("url")
        url_cell = (url, url) if url else "-"
        rows.append(
            [
                d.get("filename", ""),
                d.get("source_domain", ""),
                url_cell,
                (d.get("usability") or "-").capitalize(),
                d.get("summary") or "-",
            ]
        )

    add_table(
        doc,
        ["Filename", "Source Domain", "Source URL", "Usability", "Summary"],
        rows,
    )


# =====================================================================
# BREACH EXPOSURE
# =====================================================================


def build_breach_section(doc: DocxDocument, data: Dict[str, Any]) -> None:
    breaches_kept = data.get("breaches", {}).get("kept", [])
    if not breaches_kept:
        return

    add_section_heading(doc, "Breach Exposure", level=1)
    add_notice(doc, NOTICE_TEXT_BREACH)

    rows = []
    for b in breaches_kept:
        verdict = b.get("_llm_verdict") or {}
        rows.append(
            [
                b.get("email", ""),
                ", ".join(b.get("services", [])),
                verdict.get("exposure_summary") or "-",
            ]
        )
    add_table(doc, ["Email Address", "Services Checked", "Exposure Summary"], rows)


# =====================================================================
# DARK WEB MENTIONS
# =====================================================================


def build_darkweb_section(doc: DocxDocument, data: Dict[str, Any]) -> None:
    darkweb = data.get("darkweb", {})
    all_targets = (darkweb.get("kept", []) or []) + (darkweb.get("excluded", []) or [])
    if not all_targets:
        return

    add_section_heading(doc, "Dark Web Mentions", level=1)
    add_notice(doc, NOTICE_TEXT_DARKWEB)

    genuine_hits = [
        t
        for t in darkweb.get("kept", [])
        if (t.get("_llm_verdict") or {}).get("note") != "no_dark_web_mentions_found"
    ]

    if not genuine_hits:
        doc.add_paragraph(
            f"No dark web mentions were found across {len(all_targets)} monitored "
            f"target(s) (organization domain, company name, and named individuals)."
        )
        return

    rows = []
    for t in genuine_hits:
        verdict = t.get("_llm_verdict") or {}
        rows.append(
            [
                t.get("target", ""),
                t.get("target_type", ""),
                verdict.get("exposure_summary") or "-",
            ]
        )
    add_table(doc, ["Target", "Target Type", "Finding Summary"], rows)


# =====================================================================
# APPENDIX: DATA EXCLUSIONS & AUDIT LOG
# =====================================================================


def _excluded_rows(
    records: List[Dict[str, Any]], id_field: str, note_field: str = "note"
) -> List[List[str]]:
    rows = []
    for r in records:
        verdict = r.get("_llm_verdict") or {}
        rows.append(
            [
                r.get(id_field, "-"),
                verdict.get(note_field) or verdict.get("tier") or "-",
            ]
        )
    return rows


def build_appendix(doc: DocxDocument, data: Dict[str, Any]) -> None:
    doc.add_page_break()
    add_section_heading(doc, "Appendix: Data Exclusions & Audit Log", level=1)
    doc.add_paragraph(
        "This appendix lists items identified during data collection that were "
        "excluded from the report body above (as noise, duplicates, or "
        "unrelated false positives), and any warnings raised by the "
        "processing pipeline, for audit and transparency purposes."
    )

    stats = data.get("stats", {}) or {}
    exclusion_map = [
        (
            "domains_excluded",
            "Domains excluded",
            data.get("domains", {}).get("excluded", []),
            "domain",
        ),
        (
            "emails_excluded",
            "Emails excluded",
            data.get("emails", {}).get("excluded", []),
            "email",
        ),
        (
            "employees_excluded",
            "Employees excluded",
            data.get("employees", {}).get("excluded", []),
            "name",
        ),
        (
            "breaches_excluded",
            "Breach records excluded",
            data.get("breaches", {}).get("excluded", []),
            "email",
        ),
        (
            "darkweb_excluded",
            "Dark web targets excluded",
            data.get("darkweb", {}).get("excluded", []),
            "target",
        ),
    ]

    summary_rows = [
        [label, stats.get(key, len(records))]
        for key, label, records, _ in exclusion_map
    ]
    documents_errored = [d for d in data.get("documents", []) if d.get("error")]
    summary_rows.append(["Documents not usable", len(documents_errored)])

    add_section_heading(doc, "Exclusion Counts", level=2)
    add_table(doc, ["Category", "Count"], summary_rows)

    for _, label, records, id_field in exclusion_map:
        if not records:
            continue
        add_section_heading(doc, label, level=2)
        add_table(doc, ["Identifier", "Reason"], _excluded_rows(records, id_field))

    if documents_errored:
        add_section_heading(doc, "Documents Not Usable", level=2)
        rows = [[d.get("filename", ""), d.get("error", "-")] for d in documents_errored]
        add_table(doc, ["Filename", "Reason"], rows)

    warnings = data.get("warnings", []) or []
    add_section_heading(doc, f"Pipeline Warnings ({len(warnings)})", level=2)
    if warnings:
        for w in warnings:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(w).font.size = Pt(9)
    else:
        doc.add_paragraph("No warnings were raised during processing.")


# =====================================================================
# MAIN
# =====================================================================


def main() -> None:
    setup_logging()
    args = parse_arguments()

    company_slug = slugify_company(args.company)
    company_dir = Path("output") / company_slug
    input_file = company_dir / "llm_filtered.json"

    if not input_file.exists():
        logger.error(f"Input file not found: {input_file}. Run llm_filter.py first.")
        raise SystemExit(1)

    data = load_json(input_file)
    if not data:
        logger.error(f"{input_file} is empty or invalid.")
        raise SystemExit(1)

    report_date_display = datetime.now(timezone.utc).strftime("%B %d, %Y")
    report_date_slug = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    doc = Document()

    build_title_page(doc, data, report_date_display)
    build_summary_table(doc, data)
    build_infrastructure_section(doc, data)
    build_employees_section(doc, data)
    build_emails_section(doc, data)
    build_documents_section(doc, data)
    build_breach_section(doc, data)
    build_darkweb_section(doc, data)
    build_appendix(doc, data)

    output_file = company_dir / f"report_{company_slug}_{report_date_slug}.docx"
    doc.save(str(output_file))

    logger.info(f"Report written to: {output_file.resolve()}")
    print("\n" + "=" * 60)
    print("               REPORT GENERATION SUMMARY")
    print("=" * 60)
    print(f"Company        : {data.get('company')}")
    print(f"Output file    : {output_file.resolve()}")
    print("=" * 60 + "\n")


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="OSINT Stage: Generate the final Word report from llm_filtered.json"
    )
    parser.add_argument("--company", required=True, help="Target company name")
    return parser.parse_args()


if __name__ == "__main__":
    main()
